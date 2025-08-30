"""
Enhanced document processing system for multiple file formats.
"""
import os
import json
import hashlib
from typing import Dict, List, Optional, Tuple
from pathlib import Path
from datetime import datetime

from pypdf import PdfReader
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None
import markdown
from openai import OpenAI

class DocumentProcessor:
    def __init__(self, documents_path: str = "data/documents"):
        self.documents_path = Path(documents_path)
        self.documents_path.mkdir(parents=True, exist_ok=True)
        self.openai = OpenAI()
        
        # Supported file types
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.md': self._process_markdown,
            '.json': self._process_json
        }
    
    def process_all_documents(self) -> Dict[str, Dict]:
        """Process all documents in the documents directory."""
        processed_docs = {}
        
        for file_path in self.documents_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_types:
                try:
                    doc_info = self.process_document(str(file_path))
                    processed_docs[file_path.name] = doc_info
                except Exception as e:
                    print(f"Error processing {file_path.name}: {e}")
                    processed_docs[file_path.name] = {
                        'error': str(e),
                        'processed_at': datetime.now().isoformat()
                    }
        
        return processed_docs
    
    def process_document(self, file_path: str) -> Dict:
        """Process a single document and extract information."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_type = file_path.suffix.lower()
        
        if file_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Process the document
        content, metadata = self.supported_types[file_type](file_path)
        
        # Generate summary and keywords using AI
        summary, keywords = self._generate_ai_insights(content, file_path.name)
        
        return {
            'filename': file_path.name,
            'file_type': file_type,
            'file_size': file_path.stat().st_size,
            'content': content,
            'content_length': len(content),
            'summary': summary,
            'keywords': keywords,
            'metadata': metadata,
            'processed_at': datetime.now().isoformat(),
            'file_hash': self._get_file_hash(file_path)
        }
    
    def _process_pdf(self, file_path: Path) -> Tuple[str, Dict]:
        """Process PDF file."""
        content = ""
        metadata = {}
        
        try:
            reader = PdfReader(str(file_path))
            
            # Extract text from all pages
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    content += f"\\n--- Page {page_num} ---\\n{page_text}\\n"
            
            # Extract metadata
            if reader.metadata:
                metadata = {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'subject': reader.metadata.get('/Subject', ''),
                    'creator': reader.metadata.get('/Creator', ''),
                    'pages': len(reader.pages)
                }
        
        except Exception as e:
            raise Exception(f"Error processing PDF: {e}")
        
        return content.strip(), metadata
    
    def _process_docx(self, file_path: Path) -> Tuple[str, Dict]:
        """Process DOCX file."""
        if DocxDocument is None:
            raise Exception("python-docx not installed. Please install it to process DOCX files.")
        
        try:
            doc = DocxDocument(str(file_path))
            
            # Extract text from paragraphs
            content = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\\n"
            
            # Extract basic metadata
            properties = doc.core_properties
            metadata = {
                'title': properties.title or '',
                'author': properties.author or '',
                'subject': properties.subject or '',
                'paragraphs': len(doc.paragraphs),
                'created': properties.created.isoformat() if properties.created else '',
                'modified': properties.modified.isoformat() if properties.modified else ''
            }
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {e}")
        
        return content.strip(), metadata
    
    def _process_txt(self, file_path: Path) -> Tuple[str, Dict]:
        """Process TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = {
                'encoding': 'utf-8',
                'lines': len(content.split('\\n')),
                'words': len(content.split()),
                'characters': len(content)
            }
            
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    metadata['encoding'] = encoding
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise Exception("Could not decode file with any supported encoding")
        
        return content, metadata
    
    def _process_markdown(self, file_path: Path) -> Tuple[str, Dict]:
        """Process Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            
            # Convert to HTML for potential future use
            html_content = markdown.markdown(md_content)
            
            metadata = {
                'format': 'markdown',
                'lines': len(md_content.split('\\n')),
                'html_length': len(html_content),
                'headers': self._extract_headers(md_content)
            }
            
        except Exception as e:
            raise Exception(f"Error processing Markdown: {e}")
        
        return md_content, metadata
    
    def _process_json(self, file_path: Path) -> Tuple[str, Dict]:
        """Process JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
            
            # Convert to readable string
            content = json.dumps(json_data, indent=2)
            
            metadata = {
                'format': 'json',
                'keys': list(json_data.keys()) if isinstance(json_data, dict) else [],
                'type': type(json_data).__name__,
                'size': len(json_data) if isinstance(json_data, (list, dict)) else 1
            }
            
        except Exception as e:
            raise Exception(f"Error processing JSON: {e}")
        
        return content, metadata
    
    def _extract_headers(self, markdown_content: str) -> List[str]:
        """Extract headers from markdown content."""
        headers = []
        for line in markdown_content.split('\\n'):
            if line.strip().startswith('#'):
                headers.append(line.strip())
        return headers
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Generate MD5 hash of file."""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _generate_ai_insights(self, content: str, filename: str) -> Tuple[str, List[str]]:
        """Generate summary and keywords using AI."""
        try:
            # Limit content length for API efficiency
            content_preview = content[:3000] if len(content) > 3000 else content
            
            prompt = f"""
            Analyze the following document content from file "{filename}":

            {content_preview}

            Please provide:
            1. A concise 2-3 sentence summary
            2. 5-8 relevant keywords/topics

            Format as JSON:
            {{
                "summary": "your summary here",
                "keywords": ["keyword1", "keyword2", ...]
            }}
            """
            
            response = self.openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            
            result = json.loads(response.choices[0].message.content)
            return result.get('summary', ''), result.get('keywords', [])
            
        except Exception as e:
            print(f"Error generating AI insights: {e}")
            return f"Document: {filename}", []
    
    def search_documents(self, query: str, documents: Dict[str, Dict]) -> List[Dict]:
        """Search documents for relevant content."""
        results = []
        query_lower = query.lower()
        
        for filename, doc_info in documents.items():
            if 'error' in doc_info:
                continue
                
            score = 0
            matches = []
            
            # Search in content
            content = doc_info.get('content', '').lower()
            if query_lower in content:
                score += 3
                matches.append('content')
            
            # Search in summary
            summary = doc_info.get('summary', '').lower()
            if query_lower in summary:
                score += 2
                matches.append('summary')
            
            # Search in keywords
            keywords = doc_info.get('keywords', [])
            for keyword in keywords:
                if query_lower in keyword.lower():
                    score += 1
                    matches.append('keywords')
                    break
            
            # Search in filename
            if query_lower in filename.lower():
                score += 1
                matches.append('filename')
            
            if score > 0:
                results.append({
                    'filename': filename,
                    'score': score,
                    'matches': matches,
                    'summary': doc_info.get('summary', ''),
                    'keywords': keywords
                })
        
        return sorted(results, key=lambda x: x['score'], reverse=True)
    
    def get_document_stats(self, documents: Dict[str, Dict]) -> Dict:
        """Get statistics about processed documents."""
        total_docs = len(documents)
        successful = sum(1 for doc in documents.values() if 'error' not in doc)
        failed = total_docs - successful
        
        file_types = {}
        total_content_length = 0
        
        for doc_info in documents.values():
            if 'error' not in doc_info:
                file_type = doc_info.get('file_type', 'unknown')
                file_types[file_type] = file_types.get(file_type, 0) + 1
                total_content_length += doc_info.get('content_length', 0)
        
        return {
            'total_documents': total_docs,
            'successful_processing': successful,
            'failed_processing': failed,
            'file_types': file_types,
            'total_content_length': total_content_length,
            'average_content_length': total_content_length / successful if successful > 0 else 0
        }